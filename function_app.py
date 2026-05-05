import azure.functions as func
import json
import os
import uuid
import logging
import tempfile
import csv
from io import BytesIO, StringIO
from datetime import datetime, timezone

from openai import AzureOpenAI
from azure.storage.blob import BlobServiceClient
from azure.data.tables import TableServiceClient

import fitz
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook


app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)

SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".xlsx"]

CONTROL_DOMAINS = [
    "Security governance and assurance",
    "Policies and security standards",
    "Asset management and data classification",
    "Identity and access management",
    "Privileged access management",
    "Encryption and key management",
    "Network and infrastructure security",
    "Secure development and change management",
    "Vulnerability management and penetration testing",
    "Logging, monitoring, and incident response",
    "Backup, resilience, and disaster recovery",
    "Third party and subprocessor management",
    "Privacy, retention, and data disposal",
    "Physical and environmental security",
]


def get_env(name: str) -> str:
    value = os.getenv(name)
    if not value:
        raise ValueError(f"Missing required environment variable: {name}")
    return value


def generate_review_id() -> str:
    today = datetime.now(timezone.utc).strftime("%Y%m%d")
    short_id = uuid.uuid4().hex[:8].upper()
    return f"TPSR-{today}-{short_id}"


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def today_display() -> str:
    return datetime.now(timezone.utc).strftime("%d %B %Y")


def normalise_key(value: str) -> str:
    if not value:
        return "UNKNOWN"
    return value.strip().upper().replace(" ", "_").replace("-", "_").replace("/", "_")


def get_container_client():
    blob_service_client = BlobServiceClient.from_connection_string(get_env("TPSR_STORAGE_CONNECTION_STRING"))
    container_name = get_env("TPSR_BLOB_CONTAINER")
    container_client = blob_service_client.get_container_client(container_name)
    try:
        container_client.create_container()
    except Exception:
        pass
    return container_client


def list_supported_blobs(input_folder: str) -> list:
    container_client = get_container_client()
    prefix = input_folder.strip().rstrip("/") + "/"
    blobs = []
    for blob in container_client.list_blobs(name_starts_with=prefix):
        if any(blob.name.lower().endswith(ext) for ext in SUPPORTED_EXTENSIONS):
            blobs.append({"name": blob.name, "last_modified": str(blob.last_modified), "size": blob.size})
    if not blobs:
        raise ValueError(f"No supported evidence files found in folder: {input_folder}")
    return blobs


def download_blob_bytes(blob_path: str) -> bytes:
    return get_container_client().get_blob_client(blob_path).download_blob().readall()


def upload_bytes_to_blob(blob_path: str, data: bytes, content_type: str) -> str:
    get_container_client().get_blob_client(blob_path).upload_blob(data, overwrite=True, content_type=content_type)
    return blob_path


def upload_text_to_blob(blob_path: str, text: str, content_type: str = "text/plain") -> str:
    return upload_bytes_to_blob(blob_path, text.encode("utf-8"), content_type)


def upload_json_to_blob(blob_path: str, payload: dict) -> str:
    return upload_text_to_blob(blob_path, json.dumps(payload, indent=2), "application/json")


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    extracted_pages = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf_document:
        for page_number, page in enumerate(pdf_document, start=1):
            page_text = page.get_text("text")
            if page_text and page_text.strip():
                extracted_pages.append(f"\n--- Page {page_number} ---\n{page_text.strip()}")
    text = "\n".join(extracted_pages).strip()
    if not text:
        raise ValueError("No readable text extracted from PDF. The PDF may be scanned or image based.")
    return text


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name
    try:
        document = Document(tmp_path)
        parts = []
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table_index, table in enumerate(document.tables, start=1):
            parts.append(f"\n--- Table {table_index} ---")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        text = "\n".join(parts).strip()
        if not text:
            raise ValueError("No readable text extracted from DOCX.")
        return text
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass


def extract_text_from_xlsx_bytes(xlsx_bytes: bytes) -> str:
    workbook = load_workbook(filename=BytesIO(xlsx_bytes), data_only=True)
    parts = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        parts.append(f"\n--- Worksheet: {sheet_name} ---")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if cell is None else str(cell).strip() for cell in row]
            if any(values):
                parts.append(" | ".join(values))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("No readable text extracted from XLSX.")
    return text


def extract_text_from_csv_bytes(csv_bytes: bytes) -> str:
    raw_text = csv_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(StringIO(raw_text))
    return "\n".join(" | ".join(str(cell).strip() for cell in row) for row in reader if row).strip()


def extract_text_from_text_bytes(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()


def classify_evidence_type(blob_name: str, extracted_text: str) -> str:
    name = blob_name.lower()
    text = extracted_text.lower()[:12000]
    if "soc 2" in text or "soc2" in name or "soc 2" in name:
        return "SOC 2 report"
    if "soc 1" in text or "soc1" in name or "soc 1" in name:
        return "SOC 1 report"
    if "iso/iec 27001" in text or "iso 27001" in text or "iso27001" in name:
        return "ISO 27001 evidence"
    if "statement of applicability" in text or "soa" in name:
        return "Statement of Applicability"
    if "questionnaire" in name or "questionnaire" in text:
        return "Security questionnaire"
    if "penetration test" in text or "pentest" in name or "pen test" in name:
        return "Penetration test evidence"
    if "vulnerability" in text or "vulnerability" in name:
        return "Vulnerability management evidence"
    if "incident response" in text or "incident" in name:
        return "Incident response evidence"
    if "business continuity" in text or "disaster recovery" in text or "bcp" in name or "dr" in name:
        return "BCP or DR evidence"
    if "privacy" in text or "data protection" in text:
        return "Privacy or data protection evidence"
    if "architecture" in name or "data flow" in text or "network diagram" in text:
        return "Architecture or data flow evidence"
    if "policy" in name or "standard" in name:
        return "Security policy or standard"
    if "attestation" in name or "attest" in text:
        return "Email or management attestation"
    return "Other vendor evidence"


def extract_text_from_blob(blob_name: str) -> dict:
    file_bytes = download_blob_bytes(blob_name)
    lower_name = blob_name.lower()
    try:
        if lower_name.endswith(".pdf"):
            text = extract_text_from_pdf_bytes(file_bytes)
        elif lower_name.endswith(".docx"):
            text = extract_text_from_docx_bytes(file_bytes)
        elif lower_name.endswith(".xlsx"):
            text = extract_text_from_xlsx_bytes(file_bytes)
        elif lower_name.endswith(".csv"):
            text = extract_text_from_csv_bytes(file_bytes)
        elif lower_name.endswith(".json"):
            text = json.dumps(json.loads(extract_text_from_text_bytes(file_bytes)), indent=2)
        elif lower_name.endswith(".txt") or lower_name.endswith(".md"):
            text = extract_text_from_text_bytes(file_bytes)
        else:
            text = ""
        return {"file_name": blob_name, "evidence_type": classify_evidence_type(blob_name, text), "status": "Extracted", "error": "", "text": text}
    except Exception as ex:
        return {"file_name": blob_name, "evidence_type": "Unreadable or unsupported evidence", "status": "Failed", "error": str(ex), "text": ""}


def build_combined_evidence_text(extracted_files: list, max_chars_per_file: int = 90000) -> str:
    sections = []
    for item in extracted_files:
        sections.append(
            "\n\n" + "=" * 90 +
            f"\nFILE: {item['file_name']}\nEVIDENCE TYPE: {item['evidence_type']}\nEXTRACTION STATUS: {item['status']}\nEXTRACTION ERROR: {item['error'] or 'None'}\n" +
            "=" * 90 + "\n" + item.get("text", "")[:max_chars_per_file]
        )
    return "\n".join(sections).strip()


def trim_evidence_pack_for_ai(full_text: str, max_chars: int = 220000) -> str:
    if len(full_text) <= max_chars:
        return full_text
    priority_keywords = ["soc 2", "soc 1", "opinion", "independent service auditor", "trust services criteria", "complementary user entity controls", "subservice organization", "subservice organisation", "exceptions", "testing results", "iso 27001", "statement of applicability", "scope", "certificate", "questionnaire", "access control", "encryption", "incident response", "business continuity", "disaster recovery", "vulnerability", "penetration test", "privacy", "data protection", "retention", "subprocessor", "third party"]
    lower_text = full_text.lower()
    selected_sections = []
    for keyword in priority_keywords:
        start_search = 0
        hits = 0
        while hits < 3:
            index = lower_text.find(keyword, start_search)
            if index == -1:
                break
            start = max(index - 4000, 0)
            end = min(index + 16000, len(full_text))
            selected_sections.append(full_text[start:end])
            start_search = index + len(keyword)
            hits += 1
    if selected_sections:
        return "\n\n--- Relevant Evidence Extract ---\n\n".join(selected_sections)[:max_chars]
    return full_text[:max_chars]


def build_enterprise_tpsr_prompt(evidence_text: str, intake_context: dict, extracted_files: list) -> str:
    evidence_inventory = [{"file_name": item["file_name"], "evidence_type": item["evidence_type"], "extraction_status": item["status"], "extraction_error": item["error"]} for item in extracted_files]
    return f"""
You are a senior cyber GRC analyst performing an enterprise third party security review.

The evidence pack may include SOC 2, SOC 1, ISO 27001 certificate, Statement of Applicability, security questionnaire, security policies, technical documents, penetration test summary, privacy documents, business continuity evidence, incident response evidence, email attestation, or mixed assurance evidence.

Critical rules:
1. Use only the evidence text and intake context provided.
2. Do not invent missing facts.
3. If a fact is not available, write "Not identified in the evidence".
4. Evidence quality must be assessed before control domain reliance.
5. Questionnaire evidence is self attested unless supported by independent or operating evidence.
6. ISO 27001 certificate alone does not prove detailed operating effectiveness unless SoA or audit evidence is provided.
7. SOC 1 is not equivalent to SOC 2 for broad cyber assurance.
8. Do not ignore customer responsibilities, subprocessors, subservice organisations, or carved out dependencies.
9. Return strict valid JSON only.

Intake context:
{json.dumps(intake_context, indent=2)}

Evidence inventory:
{json.dumps(evidence_inventory, indent=2)}

Return JSON using this structure:
{{
  "review_header": {{"vendor_name": "", "service_reviewed": "", "review_type": "", "business_unit": "", "business_owner": "", "technical_owner": "", "reviewer_name": "", "review_date": "", "vendor_criticality": "", "data_sensitivity": "", "final_review_outcome": "", "residual_risk_rating": "", "next_review_date": ""}},
  "executive_summary": "",
  "vendor_and_service_context": {{"service_description": "", "organisation_use_case": "", "stores_data": "", "processes_or_transmits_data": "", "internal_access": "", "business_criticality_context": "", "integrations": "", "regulatory_obligations": ""}},
  "evidence_pack_summary": [{{"evidence_type": "", "received": "", "date_or_period": "", "scope_match": "", "assurance_value": "", "limitations": "", "source_files": []}}],
  "evidence_adequacy_and_reliance": {{"overall_evidence_reliance_rating": "", "rationale": "", "evidence_gaps": []}},
  "control_domain_assessment": [{{"control_domain": "", "evidence_reviewed": "", "analysis": "", "gap_or_limitation": "", "risk_impact": "", "rating": ""}}],
  "soc_report_analysis": {{"applicable": "", "report_type": "", "report_period": "", "service_auditor": "", "opinion_type": "", "scope_alignment": "", "exceptions": "", "cuec_summary": "", "subservice_summary": "", "reliance_conclusion": ""}},
  "iso_27001_analysis": {{"applicable": "", "certificate_status": "", "certification_body": "", "issue_and_expiry_date": "", "scope_statement": "", "service_alignment": "", "statement_of_applicability_received": "", "limitations": "", "reliance_conclusion": ""}},
  "security_questionnaire_analysis": {{"applicable": "", "questionnaire_date": "", "completed_by": "", "supporting_evidence_provided": "", "key_positive_controls": "", "key_gaps_or_unclear_responses": "", "reliance_limitation": "", "reliance_conclusion": ""}},
  "policy_and_technical_document_analysis": {{"applicable": "", "policies_and_standards": "", "access_control_documentation": "", "encryption_or_key_management_documentation": "", "architecture_or_data_flow_evidence": "", "vulnerability_or_penetration_test_evidence": "", "incident_response_evidence": "", "bcp_or_dr_evidence": "", "reliance_conclusion": ""}},
  "key_findings_limitations_and_compensating_factors": [{{"item": "", "description": "", "risk_implication": "", "compensating_factor": "", "follow_up_required": ""}}],
  "shared_responsibility_and_customer_obligations": [{{"responsibility": "", "source_evidence": "", "internal_owner": "", "evidence_required_internally": "", "risk_if_not_performed": ""}}],
  "follow_up_actions": [{{"action_required": "", "owner": "", "due_date": "", "reason": "", "evidence_expected": "", "status": ""}}],
  "residual_risk_and_review_outcome": {{"overall_evidence_reliance_rating": "", "residual_risk_rating": "", "review_outcome": "", "outcome_rationale": "", "conditions_of_approval": "", "risk_acceptance_required": "", "next_review_date": ""}},
  "approval_record": [{{"role": "", "name": "", "decision": "", "date": "", "comments": ""}}],
  "ai_extracted_vs_reviewer_selected": {{"ai_can_extract_or_draft": "", "reviewer_must_select_or_confirm": ""}},
  "quality_check": {{"evidence_types_identified": "", "evidence_quality_assessed": "", "control_domains_assessed": "", "key_limitations_marked": "", "reviewer_fields_separated": "", "outcome_supported": ""}}
}}

Control domains to assess:
{json.dumps(CONTROL_DOMAINS, indent=2)}

Evidence text:
{evidence_text}
"""


def call_ai_foundry(evidence_text: str, intake_context: dict, extracted_files: list) -> dict:
    client = AzureOpenAI(azure_endpoint=get_env("AZURE_OPENAI_ENDPOINT"), api_key=get_env("AZURE_OPENAI_KEY"), api_version=get_env("AZURE_OPENAI_API_VERSION"))
    response = client.chat.completions.create(
        model=get_env("AZURE_OPENAI_DEPLOYMENT"),
        messages=[
            {"role": "system", "content": "You are a senior cyber GRC analyst. Produce strict JSON only."},
            {"role": "user", "content": build_enterprise_tpsr_prompt(evidence_text, intake_context, extracted_files)},
        ],
        temperature=0,
        max_tokens=9000,
        response_format={"type": "json_object"},
    )
    return json.loads(response.choices[0].message.content)


def value_or_default(value, default="Not identified in the evidence") -> str:
    if value is None or value == "":
        return default
    if isinstance(value, list):
        return ", ".join(str(item) for item in value) if value else default
    return str(value)


def add_heading(document: Document, text: str, level: int = 1):
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text):
    document.add_paragraph(value_or_default(text))


def add_kv_table(document: Document, rows: list):
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    table.rows[0].cells[2].text = "Source"
    for field, value, source in rows:
        cells = table.add_row().cells
        cells[0].text = str(field)
        cells[1].text = value_or_default(value)
        cells[2].text = value_or_default(source)
    document.add_paragraph("")


def add_list_table(document: Document, items: list, headers: list, field_names: list):
    if not items:
        document.add_paragraph("Not identified in the evidence.")
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for item in items:
        cells = table.add_row().cells
        for idx, field_name in enumerate(field_names):
            cells[idx].text = value_or_default(item.get(field_name, ""))


def create_enterprise_word_report(review_id: str, ai_output: dict, extracted_files: list) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    title = document.add_heading("Enterprise Third Party Security Review", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(f"Draft report generated by Azure TPSR POC | Review ID: {review_id}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")

    header = ai_output.get("review_header", {})
    context = ai_output.get("vendor_and_service_context", {})
    reliance = ai_output.get("evidence_adequacy_and_reliance", {})
    outcome = ai_output.get("residual_risk_and_review_outcome", {})

    add_heading(document, "1. Review Header", 1)
    add_kv_table(document, [
        ("Vendor name", header.get("vendor_name"), "AI extract where available, reviewer confirm"),
        ("Service or product reviewed", header.get("service_reviewed"), "Reviewer confirm"),
        ("Review type", header.get("review_type"), "Reviewer select"),
        ("Business unit", header.get("business_unit"), "Reviewer select"),
        ("Business owner", header.get("business_owner"), "Reviewer select"),
        ("Technical owner", header.get("technical_owner"), "Reviewer select if applicable"),
        ("Reviewer", header.get("reviewer_name"), "Reviewer select"),
        ("Review date", header.get("review_date"), "Reviewer select"),
        ("Vendor criticality", header.get("vendor_criticality"), "Reviewer select"),
        ("Data sensitivity", header.get("data_sensitivity"), "Reviewer select"),
        ("Final review outcome", outcome.get("review_outcome") or header.get("final_review_outcome"), "Reviewer confirm"),
        ("Residual risk rating", outcome.get("residual_risk_rating") or header.get("residual_risk_rating"), "Reviewer confirm"),
        ("Next review date", outcome.get("next_review_date") or header.get("next_review_date"), "Reviewer confirm"),
    ])

    add_heading(document, "2. Executive Summary", 1)
    add_paragraph(document, ai_output.get("executive_summary"))

    add_heading(document, "3. Vendor and Service Context", 1)
    add_kv_table(document, [
        ("What service does the vendor provide?", context.get("service_description"), "Reviewer, AI can assist from evidence"),
        ("How does the organisation use the service?", context.get("organisation_use_case"), "Reviewer"),
        ("Does the vendor store organisation data?", context.get("stores_data"), "Reviewer confirm"),
        ("Does the vendor process or transmit organisation data?", context.get("processes_or_transmits_data"), "Reviewer confirm"),
        ("Does the vendor have access to internal systems?", context.get("internal_access"), "Reviewer confirm"),
        ("Is the service customer facing or business critical?", context.get("business_criticality_context"), "Reviewer confirm"),
        ("Are there integrations with internal systems?", context.get("integrations"), "Reviewer confirm"),
        ("Are any regulatory or contractual obligations relevant?", context.get("regulatory_obligations"), "Reviewer confirm"),
    ])

    add_heading(document, "4. Evidence Pack Summary", 1)
    add_list_table(document, ai_output.get("evidence_pack_summary", []), ["Evidence Type", "Received", "Date or Period", "Scope Match", "Assurance Value", "Limitations"], ["evidence_type", "received", "date_or_period", "scope_match", "assurance_value", "limitations"])

    add_heading(document, "5. Evidence Adequacy and Reliance Rating", 1)
    add_kv_table(document, [("Overall evidence reliance rating", reliance.get("overall_evidence_reliance_rating"), "AI draft, reviewer confirm"), ("Rationale", reliance.get("rationale"), "AI analysis"), ("Evidence gaps", value_or_default(reliance.get("evidence_gaps")), "AI analysis")])

    add_heading(document, "6. Control Domain Assessment", 1)
    add_list_table(document, ai_output.get("control_domain_assessment", []), ["Control Domain", "Evidence Reviewed", "Analysis", "Gap or Limitation", "Risk Impact", "Rating"], ["control_domain", "evidence_reviewed", "analysis", "gap_or_limitation", "risk_impact", "rating"])

    add_heading(document, "7. Evidence Specific Analysis Blocks", 1)
    for title_text, key, rows in [
        ("7.1 SOC Report Analysis", "soc_report_analysis", ["applicable", "report_type", "report_period", "service_auditor", "opinion_type", "scope_alignment", "exceptions", "cuec_summary", "subservice_summary", "reliance_conclusion"]),
        ("7.2 ISO 27001 Analysis", "iso_27001_analysis", ["applicable", "certificate_status", "certification_body", "issue_and_expiry_date", "scope_statement", "service_alignment", "statement_of_applicability_received", "limitations", "reliance_conclusion"]),
        ("7.3 Security Questionnaire Analysis", "security_questionnaire_analysis", ["applicable", "questionnaire_date", "completed_by", "supporting_evidence_provided", "key_positive_controls", "key_gaps_or_unclear_responses", "reliance_limitation", "reliance_conclusion"]),
        ("7.4 Policy and Technical Document Analysis", "policy_and_technical_document_analysis", ["applicable", "policies_and_standards", "access_control_documentation", "encryption_or_key_management_documentation", "architecture_or_data_flow_evidence", "vulnerability_or_penetration_test_evidence", "incident_response_evidence", "bcp_or_dr_evidence", "reliance_conclusion"]),
    ]:
        add_heading(document, title_text, 2)
        block = ai_output.get(key, {})
        add_kv_table(document, [(field.replace("_", " ").title(), block.get(field), "AI analysis") for field in rows])

    add_heading(document, "8. Key Findings, Limitations, and Compensating Factors", 1)
    add_list_table(document, ai_output.get("key_findings_limitations_and_compensating_factors", []), ["Item", "Description", "Risk Implication", "Compensating Factor", "Follow Up Required"], ["item", "description", "risk_implication", "compensating_factor", "follow_up_required"])

    add_heading(document, "9. Shared Responsibility and Customer Obligations", 1)
    add_list_table(document, ai_output.get("shared_responsibility_and_customer_obligations", []), ["Responsibility", "Source Evidence", "Internal Owner", "Evidence Required Internally", "Risk If Not Performed"], ["responsibility", "source_evidence", "internal_owner", "evidence_required_internally", "risk_if_not_performed"])

    add_heading(document, "10. Follow Up Actions", 1)
    add_list_table(document, ai_output.get("follow_up_actions", []), ["Action Required", "Owner", "Due Date", "Reason", "Evidence Expected", "Status"], ["action_required", "owner", "due_date", "reason", "evidence_expected", "status"])

    add_heading(document, "11. Residual Risk and Review Outcome", 1)
    add_kv_table(document, [("Overall evidence reliance rating", outcome.get("overall_evidence_reliance_rating"), "AI draft, reviewer confirm"), ("Residual risk rating", outcome.get("residual_risk_rating"), "Reviewer confirm"), ("Review outcome", outcome.get("review_outcome"), "Reviewer confirm"), ("Outcome rationale", outcome.get("outcome_rationale"), "AI draft"), ("Conditions of approval", outcome.get("conditions_of_approval"), "Reviewer confirm"), ("Risk acceptance required", outcome.get("risk_acceptance_required"), "Reviewer confirm"), ("Next review date", outcome.get("next_review_date"), "Reviewer confirm")])

    add_heading(document, "Appendix A. Evidence Files Processed", 1)
    file_rows = [{"file_name": item.get("file_name"), "evidence_type": item.get("evidence_type"), "status": item.get("status"), "error": item.get("error") or "None"} for item in extracted_files]
    add_list_table(document, file_rows, ["File Name", "Evidence Type", "Extraction Status", "Error"], ["file_name", "evidence_type", "status", "error"])

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_path = tmp.name
    try:
        document.save(temp_path)
        with open(temp_path, "rb") as f:
            return f.read()
    finally:
        try:
            os.remove(temp_path)
        except Exception:
            pass


def save_review_to_table(review_id: str, input_folder: str, ai_output: dict, extracted_text_blob_path: str, ai_output_blob_path: str, word_report_blob_path: str, evidence_inventory_blob_path: str, status: str, error_message: str = ""):
    table_service_client = TableServiceClient.from_connection_string(get_env("TPSR_STORAGE_CONNECTION_STRING"))
    table_name = get_env("TPSR_TABLE_NAME")
    try:
        table_service_client.create_table(table_name)
    except Exception:
        pass
    table_client = table_service_client.get_table_client(table_name)
    header = ai_output.get("review_header", {})
    outcome = ai_output.get("residual_risk_and_review_outcome", {})
    reliance = ai_output.get("evidence_adequacy_and_reliance", {})
    review_type = value_or_default(header.get("review_type"), "TPSR")
    entity = {
        "PartitionKey": normalise_key(review_type),
        "RowKey": review_id,
        "reviewId": review_id,
        "vendorName": value_or_default(header.get("vendor_name")),
        "serviceReviewed": value_or_default(header.get("service_reviewed")),
        "reviewType": review_type,
        "businessUnit": value_or_default(header.get("business_unit")),
        "businessOwner": value_or_default(header.get("business_owner")),
        "technicalOwner": value_or_default(header.get("technical_owner")),
        "reviewerName": value_or_default(header.get("reviewer_name")),
        "vendorCriticality": value_or_default(header.get("vendor_criticality")),
        "dataSensitivity": value_or_default(header.get("data_sensitivity")),
        "status": status,
        "inputFolder": input_folder,
        "overallEvidenceRelianceRating": value_or_default(reliance.get("overall_evidence_reliance_rating") or outcome.get("overall_evidence_reliance_rating")),
        "residualRiskRating": value_or_default(outcome.get("residual_risk_rating")),
        "finalReviewOutcome": value_or_default(outcome.get("review_outcome")),
        "riskAcceptanceRequired": value_or_default(outcome.get("risk_acceptance_required")),
        "nextReviewDate": value_or_default(outcome.get("next_review_date")),
        "extractedTextBlobPath": extracted_text_blob_path,
        "aiOutputBlobPath": ai_output_blob_path,
        "wordReportBlobPath": word_report_blob_path,
        "evidenceInventoryBlobPath": evidence_inventory_blob_path,
        "createdUtc": utc_now_iso(),
        "modelUsed": get_env("AZURE_OPENAI_DEPLOYMENT"),
        "errorMessage": error_message,
    }
    table_client.upsert_entity(entity)


@app.route(route="tpsr_review", methods=["POST"])
def tpsr_review(req: func.HttpRequest) -> func.HttpResponse:
    logging.info("Enterprise TPSR evidence pack review request received")
    review_id = generate_review_id()
    try:
        body = req.get_json()
        source_type = body.get("sourceType", "evidencePack").strip()
        input_folder = body.get("inputFolder", "").strip()
        if source_type != "evidencePack":
            return func.HttpResponse(json.dumps({"status": "failed", "message": "This function expects sourceType = evidencePack."}), status_code=400, mimetype="application/json")
        if not input_folder:
            return func.HttpResponse(json.dumps({"status": "failed", "message": "inputFolder is required. Example: demo-inputs/software-vendor-pack"}), status_code=400, mimetype="application/json")

        intake_context = {
            "vendorName": body.get("vendorName", "").strip(),
            "serviceReviewed": body.get("serviceReviewed", "").strip(),
            "reviewType": body.get("reviewType", "Initial Review").strip(),
            "businessUnit": body.get("businessUnit", "To be confirmed").strip(),
            "businessOwner": body.get("businessOwner", "To be confirmed").strip(),
            "technicalOwner": body.get("technicalOwner", "To be confirmed").strip(),
            "reviewerName": body.get("reviewerName", "To be confirmed").strip(),
            "reviewDate": body.get("reviewDate", today_display()).strip(),
            "vendorCriticality": body.get("vendorCriticality", "To be confirmed").strip(),
            "dataSensitivity": body.get("dataSensitivity", "To be confirmed").strip(),
            "organisationUseCase": body.get("organisationUseCase", "To be confirmed by reviewer").strip(),
            "nextReviewDate": body.get("nextReviewDate", "To be confirmed").strip(),
        }

        supported_blobs = list_supported_blobs(input_folder)
        extracted_files = [extract_text_from_blob(blob["name"]) for blob in supported_blobs]
        combined_evidence_text = build_combined_evidence_text(extracted_files)
        ai_ready_text = trim_evidence_pack_for_ai(combined_evidence_text)

        extracted_text_blob_path = f"{review_id}/extracted-evidence.txt"
        evidence_inventory_blob_path = f"{review_id}/evidence-inventory.json"
        ai_output_blob_path = f"{review_id}/ai-output.json"
        word_report_blob_path = f"{review_id}/draft-enterprise-tpsr-review.docx"

        upload_text_to_blob(extracted_text_blob_path, combined_evidence_text)
        upload_json_to_blob(evidence_inventory_blob_path, {"reviewId": review_id, "inputFolder": input_folder, "files": [{"file_name": item["file_name"], "evidence_type": item["evidence_type"], "status": item["status"], "error": item["error"]} for item in extracted_files]})

        ai_output = call_ai_foundry(ai_ready_text, intake_context, extracted_files)
        header = ai_output.setdefault("review_header", {})
        for ai_key, intake_key in [("review_type", "reviewType"), ("business_unit", "businessUnit"), ("business_owner", "businessOwner"), ("technical_owner", "technicalOwner"), ("reviewer_name", "reviewerName"), ("review_date", "reviewDate"), ("vendor_criticality", "vendorCriticality"), ("data_sensitivity", "dataSensitivity"), ("next_review_date", "nextReviewDate")]:
            if not header.get(ai_key):
                header[ai_key] = intake_context[intake_key]

        upload_json_to_blob(ai_output_blob_path, ai_output)
        word_report_bytes = create_enterprise_word_report(review_id, ai_output, extracted_files)
        upload_bytes_to_blob(word_report_blob_path, word_report_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        save_review_to_table(review_id, input_folder, ai_output, extracted_text_blob_path, ai_output_blob_path, word_report_blob_path, evidence_inventory_blob_path, "Completed")

        outcome = ai_output.get("residual_risk_and_review_outcome", {})
        reliance = ai_output.get("evidence_adequacy_and_reliance", {})
        response_payload = {
            "reviewId": review_id,
            "status": "Completed",
            "vendorName": value_or_default(header.get("vendor_name")),
            "serviceReviewed": value_or_default(header.get("service_reviewed")),
            "overallEvidenceRelianceRating": value_or_default(reliance.get("overall_evidence_reliance_rating") or outcome.get("overall_evidence_reliance_rating")),
            "residualRiskRating": value_or_default(outcome.get("residual_risk_rating")),
            "draftReviewOutcome": value_or_default(outcome.get("review_outcome")),
            "storage": {
                "blobContainer": get_env("TPSR_BLOB_CONTAINER"),
                "inputFolder": input_folder,
                "extractedTextBlobPath": extracted_text_blob_path,
                "evidenceInventoryBlobPath": evidence_inventory_blob_path,
                "aiOutputBlobPath": ai_output_blob_path,
                "wordReportBlobPath": word_report_blob_path,
                "tableName": get_env("TPSR_TABLE_NAME"),
                "partitionKey": normalise_key(value_or_default(header.get("review_type"), "TPSR")),
                "rowKey": review_id,
            },
            "filesProcessed": [{"fileName": item["file_name"], "evidenceType": item["evidence_type"], "status": item["status"], "error": item["error"]} for item in extracted_files],
            "aiOutput": ai_output,
        }
        return func.HttpResponse(json.dumps(response_payload, indent=2), status_code=200, mimetype="application/json")
    except Exception as ex:
        logging.exception("Enterprise TPSR review failed")
        return func.HttpResponse(json.dumps({"reviewId": review_id, "status": "Failed", "error": str(ex)}, indent=2), status_code=500, mimetype="application/json")
